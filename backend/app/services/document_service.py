import csv
import io
import json
from pathlib import Path

from app.detectors.hybrid import detect_pii
from app.redaction.service import redact_text


def extract_text(filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".log"}:
        return content.decode("utf-8", errors="replace"), "text/plain"
    if suffix == ".json":
        value = json.loads(content.decode("utf-8"))
        return json.dumps(value, ensure_ascii=False, indent=2), "application/json"
    if suffix == ".csv":
        rows = csv.reader(io.StringIO(content.decode("utf-8", errors="replace")))
        return "\n".join(",".join(row) for row in rows), "text/csv"
    if suffix == ".pdf":
        import fitz
        with fitz.open(stream=content, filetype="pdf") as document:
            return "\n".join(page.get_text() for page in document), "application/pdf"
    if suffix == ".docx":
        from docx import Document
        document = Document(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, JSON, or CSV.")


def redact_document_file(filename: str, content: bytes, masking_mode: str) -> tuple[bytes, str, str]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        import fitz
        document = fitz.open(stream=content, filetype="pdf")
        for page in document:
            page_text = page.get_text()
            entities = detect_pii(page_text)
            for entity in entities:
                value = page_text[entity.start:entity.end]
                rectangles = page.search_for(value)
                if not rectangles:
                    continue
                local_entity = entity.model_copy(update={"start": 0, "end": len(value)})
                replacement = redact_text(value, [local_entity], masking_mode)
                for rectangle in rectangles:
                    if masking_mode == "black":
                        page.add_redact_annot(rectangle, fill=(0, 0, 0))
                    else:
                        font_size = max(5, min(10, rectangle.width / max(len(replacement), 1) * 1.7))
                        page.add_redact_annot(rectangle, text=replacement, fontname="helv", fontsize=font_size, fill=(1, 1, 1), text_color=(0, 0, 0), align=0)
            page.apply_redactions()
        output = document.tobytes(garbage=4, deflate=True)
        document.close()
        return output, "application/pdf", filename
    if suffix == ".docx":
        from docx import Document
        document = Document(io.BytesIO(content))
        for paragraph in list(document.paragraphs) + [paragraph for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]:
            entities = detect_pii(paragraph.text)
            paragraph.text = redact_text(paragraph.text, entities, masking_mode)
        output = io.BytesIO()
        document.save(output)
        return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename
    text, media_type = extract_text(filename, content)
    entities = detect_pii(text)
    return redact_text(text, entities, masking_mode).encode("utf-8"), media_type, filename
